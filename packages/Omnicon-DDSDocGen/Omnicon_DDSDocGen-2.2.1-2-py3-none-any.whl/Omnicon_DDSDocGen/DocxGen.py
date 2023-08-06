import inspect
import os
import shutil

import docx
import win32com.client
from docx.enum.table import WD_TABLE_ALIGNMENT

from .DocGenInterface import DocGenInterface
from docx.oxml.ns import qn
from docx.oxml import CT_R, OxmlElement

from . import Logger

from docx.table import Table
from docx.text.paragraph import Paragraph


class DocxGen(DocGenInterface):

    def __init__(self, title: str, version: str, time: str):
        super().__init__(title, version, time)
        self.logger = Logger.add_logger(__name__)

        # Create a new '.docx' file based on the styling info from the template
        self.document = None
        self.doc_template_path = DocxGen.get_doc_template_path()
        if not self.doc_template_path:
            self.logger.fatal("Template.docx wasn't found. Cannot create doc")
            raise Exception("Template.docx wasn't found. Cannot create doc")
        else:
            self.logger.debug("doc template path at " + self.doc_template_path)
            self.document = docx.Document(self.doc_template_path)

        # Clear the document
        self.document._body.clear_content()
        self.add_doc_title_page()
        self.add_toc_page()

    @staticmethod
    def get_doc_template_path() -> str:
        doc_name = "Template.docx"
        local_path = os.path.join(os.path.dirname(os.path.realpath(__file__)), doc_name)
        cwd_path = os.path.join(os.path.join(os.getcwd(), doc_name))
        if os.path.exists(local_path):
            return local_path
        elif os.path.exists(cwd_path):
            return cwd_path
        else:
            # https://importlib-resources.readthedocs.io/en/latest/using.html
            from importlib_resources import files
            whl_path = files('').joinpath(doc_name)
            if whl_path.is_file():
                return str(whl_path)
        return None

    def add_doc_title_page(self):
        """
        This function adds a title page, with creation time, including header & footer,
        to be appended as a preamble section.
        """
        self.logger.debug(self.add_doc_title_page.__name__)
        # Add the title:
        self.document.add_heading(self.title, 0)

        par: Paragraph = self.document.add_paragraph('Version ' + self.version)
        par.style = self.document.styles['version']
        par: Paragraph = self.document.add_paragraph('Generated on ' + self.time)
        par.style = self.document.styles['Subtitle']

    def add_toc_page(self):
        """
        This function adds a table of contents with hyperlink.
        """
        self.logger.debug(self.add_toc_page.__name__)
        self.add_new_page()
        # https://stackoverflow.com/questions/18595864/python-create-a-table-of-contents-with-python-docx-lxml
        par = self.document.add_paragraph("Table of Content", "Table of Content")
        run = par.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-2" \\h \\z \\u'  # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)

    def add_chapter(self, section_title, sub_section_title=""):
        """
        This function adds a new chapter, on a new page.
        :param section_title: section name.
        :param sub_section_title: sub section name.
        """
        self.logger.debug(self.add_chapter.__name__)
        self.document.add_heading(section_title, 1)

        if sub_section_title != "":
            sub_section_par = self.document.add_paragraph()
            sub_section_par.add_run('Type').underline = True
            sub_section_par.add_run(": " + sub_section_title)

    @staticmethod
    def set_repeat_table_header(row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    def add_table_header(self, listOfTitles, bLongHeader=True, color="grey"):
        """
        This function creates a table, and adds a table header with the titles.

        :param listOfTitles: the titles of the columns at the table header.
        :param bLongHeader: not in use, inherited from base class. May be needed in
                            other implementation environments. Defaults to True.
        :param color: the color of the table header, for better visualization.
                      Defaults to grey.
        """
        self.logger.debug(self.add_table_header.__name__)
        table: Table = self.document.add_table(rows=1, cols=len(listOfTitles), style='Table Grid')
        table.style = self.document.styles['Table Grid']
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        table_header_cells = table.rows[0].cells
        for i, title in enumerate(listOfTitles):
            table_header_cells[i].text = title
            table_header_cells[i].paragraphs[0].runs[0].font.bold = True
        DocxGen.set_repeat_table_header(table.rows[0])
        return table

    def add_table_row(self, theTable, cells_text, align='c'):  # align=centered
        """
        This function adds a row for a table.
        :param theTable: the table object.
        :param cells_text: the data for the columns - as text.
        :param align:  alignment - centered (not in use here)
        """
        self.logger.debug(self.add_table_row.__name__)
        row_cells = theTable.add_row().cells
        for i, cell_content in enumerate(cells_text):
            row_cells[i].text += cell_content

        if cells_text[0] == "":
            # When this element is on the basic level of the type, write the name in bold.
            row_cells[1].paragraphs[0].runs[0].font.bold = True

    def add_new_page(self):
        """
        This function adds a new page.
        """
        self.logger.debug(self.add_new_page.__name__)
        self.document.add_page_break()

    def add_section(self):
        """
        This function adds a new section. Not implemented.
        """
        self.logger.debug(self.add_section.__name__)
        pass

    def add_description(self, descr):
        """
        This function adds a description (within a new section/subsection)..
        """
        self.logger.debug(self.add_description.__name__)
        description = self.document.add_paragraph()
        description.add_run('Description').underline = True
        description.add_run(": " + descr)
        self.add_new_line()

    def add_new_line(self):
        """
        This function adds a new line.
        """
        self.logger.debug(self.add_new_line.__name__)
        # Add an empty line
        self.document.add_paragraph("")

    def generate_doc(self, output_file_name: str, temp_folder: str):
        """
        This function invokes the generation of the docx file then saves it in the requested folder.
        """
        self.logger.debug(self.generate_doc.__name__)
        output_file_name = f'{output_file_name}.docx'
        temp_output_path = os.path.join(temp_folder, output_file_name)
        self.logger.info(f"Generating {output_file_name}")

        try:
            self.document.save(temp_output_path)
            self.logger.debug(f"Initial writing to temporary folder succeeded.")
        except Exception as err:
            self.logger.error(f"The operation of saving into temporary folder has FAILED", exc_info=True)
            return

        # Update TOC
        try:
            self.update_toc(temp_output_path)
            self.logger.debug("File was saved at temporary folder")
        except Exception as err:
            self.logger.warn("Office is not installed. Cannot update table of content. User will have to perform "
                             "manually (F9 after opening document). Error = {}".format(str(err)))

    def finalize_doc(self):
        self.logger.debug(self.finalize_doc.__name__)
        pass
        # TODO set_autofit takes a long time; Could take minutes in large documents. Need a better solution
        # self.set_autofit()

    def update_toc(self, temp_path):
        self.logger.debug(self.update_toc.__name__)
        # automatically updates the table table of content
        # https://github.com/python-openxml/python-docx/issues/36
        file_path = os.path.join(os.path.join(os.getcwd(), temp_path))
        word = win32com.client.DispatchEx("Word.Application")
        doc = word.Documents.Open(file_path)
        doc.TablesOfContents(1).Update()
        doc.Close(SaveChanges=True)
        word.Quit()

    def set_autofit(self):
        """
        Make all table autofit to content
        """
        self.logger.debug(self.set_autofit.__name__)
        # TODO check if can set auto fit to window
        # https://github.com/python-openxml/python-docx/issues/209
        for t_idx, table in enumerate(self.document.tables):
            self.document.tables[t_idx].autofit = True
            self.document.tables[t_idx].allow_autofit = True
            self.document.tables[t_idx]._tblPr.xpath("./w:tblW")[0].attrib[
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"] = "auto"
            for row_idx, r_val in enumerate(self.document.tables[t_idx].rows):
                for cell_idx, c_val in enumerate(self.document.tables[t_idx].rows[row_idx].cells):
                    self.document.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.type = 'auto'
                    self.document.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.w = 0

# https://github.com/python-openxml/python-docx/commit/65db85311e9de6e50add607be169e57f8fcc7591
#         cp = self.document.paragraphs[0]
#         print(cp.text)
#         new_paragraph = cp.insert_paragraph_before('barfoo')


# TODO change table styling
